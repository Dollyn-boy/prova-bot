from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Pt
import sys

def get_mark(doc):
    for paragrafo in doc.paragraphs:
        if paragrafo.text == "[MARK]":
            return paragrafo
    raise ValueError("Template de documento word inválido, por favor adicionar '[MARK]' onde deseja começar")

def gerar_arquivo(questoes):
    doc = Document("default-ana-nete.docx")
    estilo_padrao = doc.styles['Normal']
    paragrafo_estilo = estilo_padrao.paragraph_format
    paragrafo_estilo.line_spacing = Pt(12)  # Ajuste o espaçamento conforme necessário
    


    for i, questao in enumerate(questoes):
        numero_da_questao = f"{i + 1}"
        p_enunciado = doc.add_paragraph()
        p_enunciado.style = 'Normal'

        # Adiciona o número da questão em negrito
        run_numero_questao = p_enunciado.add_run(numero_da_questao)
        font = run_numero_questao.font
        font.bold = True

        # Adiciona o enunciado da questão
        p_enunciado.add_run(f'. {questao["enunciado"]}')

        #Adiciona Alternativas
        numero_de_alternativas = len(questao) - 1
        for j in range(numero_de_alternativas):
            alternativa_key = f'alternativa{j}'
            alternativa = questao[alternativa_key]
            doc.add_paragraph(f"{chr(j+97)}) {alternativa}", style='Normal')  # You can choose a different style if needed

        doc.add_paragraph("\n")
    
    doc.save("prova_gerada.docx")



def get_enunciado(raw_questao):
    try:
        raw_enunciado = WebDriverWait(raw_questao, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, "q-question-enunciation"))
        )
        return raw_enunciado.get_attribute("aria-label")
    
    except Exception as e:
        raise ValueError(f"Error: {e}")

def get_alternativas(raw_questao):
    try:
        raw_alternativas = WebDriverWait(raw_questao, 10).until(
            EC.presence_of_all_elements_located((By.CLASS_NAME, "js-alternative-content"))
        )
        return [alternativa.text for alternativa in raw_alternativas if alternativa.text]

    except Exception as e:
        raise ValueError(f"Error: {e}")


def get_user_input():
    return sys.argv[1].replace("_", " ")

def raspar_dados(assunto, driver):
    driver.get("https://google.com")

    # Aguarda a barra de pesquisa estar presente e a retorna 
    search_bar = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.CLASS_NAME, "gLFyf"))
    )

    # Realiza a pesquisa no Google
    search_bar.clear()
    search_bar.send_keys(f"Questões História {assunto}")
    search_bar.submit()

    # Aguarda a presença do link do Qconcursos e retorna
    qconcursos_link = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.PARTIAL_LINK_TEXT, "Qconcursos"))
    )

    qconcursos_link.click()

    # Aguarda a presença de elementos com a classe "q-question-enunciation" e ""
    raw_questoes = WebDriverWait(driver, 10).until(
        EC.presence_of_all_elements_located((By.CLASS_NAME, "q-question-item"))
    )

    return raw_questoes

    

def tratar_questao(raw_questoes):
    questoes = []

    for raw_questao in raw_questoes:
        enunciado = get_enunciado(raw_questao)
        alternativas = get_alternativas(raw_questao)

        questao = {"enunciado": enunciado.strip()}

        for i, alternativa in enumerate(alternativas):
            questao[f"alternativa{i}"] = alternativa.strip()

        questoes.append(questao)

    return questoes

def ativar_driver():
    service = Service(executable_path="chromedriver.exe")
    driver = webdriver.Chrome(service=service)
    return driver

def desativar_driver(driver):
    driver.quit()

def main():
    try:
        assunto = get_user_input()
        driver = ativar_driver()
        raw_questoes = raspar_dados(assunto, driver)
        questoes = tratar_questao(raw_questoes)
        gerar_arquivo(questoes)
        desativar_driver(driver)
    except ValueError as e:
        print(f"ERROR: {e}")
        desativar_driver(driver)
        sys.exit(1)

if __name__ == "__main__":
    main()
